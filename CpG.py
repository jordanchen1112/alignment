import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import random
# from scipy.stats import norm

def randomATCG(n):
	list1 = ['A','T','C','G']
	string = ''
	for i in range(n):
		b=random.randint(0,len(list1)-1)
		string += list1[b]

	return string

def is_atcg(seq:str):
	# 判斷seq內是否只含有'A','T','C','G'
	for n in seq:
		if n not in ['A','T','C','G']:
			return False
	return True

def atcg_matrix():
	# Create empty ATCG dict (ex: {'A':{'A'}:0} )
	# 用來算例如 A->T 有幾個
	d = {}
	for n1 in ['A','T','C','G']:
		d2={}
		for n2 in ['A','T','C','G']:		
			d2[n2] = 0
		d[n1] = d2

	return d

def counts(model, seq):	
	# 給一個seq, 計算此seq中 ex:A->T 這類有幾個
	if is_atcg(seq):
		for i in range(len(seq)-1):
			model[seq[i]][seq[i+1]] += 1
			i += 1
	else:
		print(f'ATCG Error: {seq}')
		return 0
	return model

def trun_to_p(model):
	# 將個數轉成機率
	for a1 in list(model.keys()):
		s = 0
		for a2 in list(model[a1].keys()):
			s += model[a1][a2]
		if s != 0:
			for a2 in list(model[a1].keys()):
				model[a1][a2] = round(model[a1][a2] / s, 15)
		else:
			print('ZeroDivisionError')
	return model

def read_word(word_name):
	# 讀word檔, 得到+/-的seqs (in list)
	positive_seqs = []
	negative_seqs = []
	name = word_name
	path = ''
	doc = docx.Document(path + name + '.docx')
	for i in range(len(doc.paragraphs)):
		if '>' and '+' in doc.paragraphs[i].text:
			if is_atcg(doc.paragraphs[i+1].text):
				positive_seqs.append(doc.paragraphs[i+1].text)
		elif '>' and '-' in doc.paragraphs[i].text:
			if is_atcg((doc.paragraphs[i+1].text)):
				negative_seqs.append(doc.paragraphs[i+1].text)

	return positive_seqs, negative_seqs

def create_model(positive_seqs, negative_seqs):
	# Create +/- transition probabilities fro each model

	# Create empty dict of transition posiibilities for positive and megative model
	model_positive = atcg_matrix()
	model_negative = atcg_matrix()

	# Record Counts of each pair in the given model
	for seq in positive_seqs:
		counts(model_positive,seq)
	for seq in negative_seqs:
		counts(model_negative, seq)

	# Turn the counts in the model to transition possibilities 
	model_positive = trun_to_p(model_positive)
	model_negative = trun_to_p(model_negative)

	print(pd.DataFrame(model_positive).T,'\n')
	print(pd.DataFrame(model_negative).T)

	return model_positive, model_negative

def Beta(model_positive, model_negative):
	# Calculate beta (log likelihood ratios) by log2(division)
	beta = atcg_matrix()
	for n1 in ['A','T','C','G']:
		for n2 in ['A','T','C','G']:
			beta[n1][n2] = np.log2(model_positive[n1][n2] / model_negative[n1][n2])
	print(pd.DataFrame(beta).T)
	return beta

def Bits(seq:str = None, seq_list:list = None, beta = None):
	# Calulate bits of seq by beta
	# imput can be either single seq or seqs list
	if seq:
		# test = atcg_matrix()
		bit = 0
		if is_atcg(seq):
			for i in range(len(seq)-1):
				bit += beta[seq[i]][seq[i+1]]
				# test[seq[i]][seq[i+1]] += 1
			bit = bit / len(seq)
			# print(pd.DataFrame(test).T)
		else:
			print('ATCG error')
			return 0
	elif seq_list:
		bit = []
		for seq in seq_list:
			b = 0
			if not is_atcg(seq):
				print('ATCG error')
				continue
			for i in range(len(seq)-1):
				b += beta[seq[i]][seq[i+1]]
			bit.append(b / len(seq))
	return bit

def kl(name, positive_bits, negative_bits):

	# Calculate the KL divegence of each value
	seq_x = 'CCTCGTGGCCTCAGGCCTCTTGACTTCAGGCGGTTCTGTTTAATCAAGTGACATCTTCCCGAGGCTCCCTG'

	# weights for density hist plots
	weights1 = np.ones_like(positive_bits)/float(len(positive_bits))
	weights2 = np.ones_like(negative_bits)/float(len(negative_bits))

	# KL
	h1, b1 = np.histogram(positive_bits, weights=weights1,bins = 100,range=(-1,1))
	h2, b2 = np.histogram(negative_bits, weights = weights2, bins = 100,range=(-1,1))
	# D(+||-)
	h3 = h1 - h2 
	# D(-||+)
	h4= h2 - h1 
	# Adjudt the shape of h3,4 and b1
	h3 = np.append(h3,[0])
	h4 = np.append(h4,[0])

	plt.hist(positive_bits, weights=weights1,color = 'black', alpha = 0.4)
	plt.hist(negative_bits, weights=weights2,color = 'black', alpha = 0.2)
	plt.plot(b1,h3, color = 'blue', alpha = 0.5, label = 'D(+||-)')
	plt.fill_between(b1, 0, h3, facecolor='blue', alpha=0.5)
	plt.plot(b1,h4, color = 'red', alpha = 0.5, label = 'D(-||+)')
	plt.fill_between(b1, 0, h4, facecolor='red', alpha=0.5)
	plt.xlabel('Bits')
	plt.ylabel('Probabilities')
	plt.title('KL area')
	plt.legend()
	plt.show()

def ROC(numbers, positive_bits, negative_bits):
	# Numbers means how many values(dots) to plot on the ROC curve
	i = 0
	xlist = []
	ylist = []
	x = min(negative_bits)
	Youden = -1
	cutoff = -1
	while i < numbers:
		TP = len([n for n in positive_bits if n >= x])
		FN = len([n for n in positive_bits if n < x])
		TN = len([n for n in negative_bits if n <= x])
		FP = len([n for n in negative_bits if n > x])
		sensitivity = TP / (TP + FN) 
		specificity = TN / (TN + FP)
		xlist.append(1 - specificity)
		ylist.append(sensitivity)
		if sensitivity + specificity - 1 > Youden:
			Youden = sensitivity - (1 - specificity)
			cutoff = x
			cutpoint = [(1 - specificity),sensitivity]  
		x += (max(positive_bits) - min(negative_bits)) / numbers
		i += 1

	plt.scatter(cutpoint[0],cutpoint[1], color = 'black', s = 100)
	plt.plot(xlist, ylist, color = 'firebrick',linewidth= 3)
	plt.fill_between(xlist, 0, ylist, facecolor='red', alpha=0.1)
	plt.xlabel('False-Positive Ratio (1-specificity)')
	plt.ylabel('True-Positive Ratio (sensitivity)')
	plt.title(f'ROC curve (cutoff = {round(cutoff,3)})')
	plt.show()
	
	return cutoff, Youden

if __name__ == '__main__':
	# 計算+/-models and beta matrix, 將train data丟進去beta matrix
	name = 'CpG_island_sequences_2022'
	positive_seqs, negative_seqs = read_word(word_name = name)
	model_positive, model_negative = create_model(positive_seqs, negative_seqs)
	beta = Beta(model_positive, model_negative)
	positive_bits = Bits(seq_list = positive_seqs, beta = beta)
	negative_bits = Bits(seq_list = negative_seqs, beta = beta)

	seq_x = 'CCTCGTGGCCTCAGGCCTCTTGACTTCAGGCGGTTCTGTTTAATCAAGTGACATCTTCCCGAGGCTCCCTG'
	x_bit = Bits(seq = seq_x, beta = beta)
	cutoff, Youden = ROC(100, positive_bits, negative_bits)

	plt.hist(positive_bits,color = 'black', alpha = 0.8, label = 'CpG +')
	plt.hist(negative_bits,color = 'black', alpha = 0.5, label = 'CpG -')
	plt.xlabel('Bits')
	plt.ylabel('Counts')
	plt.vlines(cutoff,0,5,color="black",label ='Cut off')
	plt.vlines(x_bit,0,5,color="red", label ='X bit')
	plt.legend()
	plt.show()

	kl(name, positive_bits, negative_bits)

	print('\nResult\n-------------------------------')
	print(f'Bits of seq = {x_bit}\nCutoff = {cutoff}')
	if x_bit < cutoff:
		print('Bits of seq is smaller than Cutoff, hence the sequnce belongs to non-CpG island.\n-------------------------------')
	else:
		print('Bits of seq is larger than Cutoff, hence the sequnce belongs to CpG island.\n-------------------------------')
	